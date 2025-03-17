import json
import os
from datetime import datetime
import uuid
import threading
from docx import Document
import re
import openai
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

client = OpenAI(api_key = os.getenv("OPENAI_API_KEY"))

class Core:
    def __init__(self):
        self.contract_directory = "../store/json"
        self.contract_docx_directory = "../store/docx"
        self.lock = threading.Lock() # lock for thread safety
        if not os.path.exists(self.contract_directory):
            os.makedirs(self.contract_directory)
        if not os.path.exists(self.contract_docx_directory):
            os.makedirs(self.contract_docx_directory)

    def _generate_id(self):
        return str(uuid.uuid4())

    def _get_contract_path(self, contract_id):
        return f"{self.contract_directory}/{contract_id}.json"

    def create_contract(self, creator_id, creator_name, title, description, collaborators=None):
        contract_id = self._generate_id()
        creation_date = datetime.now().isoformat()
        
        contract = {
            "metadata": {
                "contract_id": contract_id,
                "creator_id": creator_id,
                "creator_name": creator_name,
                "title": title,
                "description": description,
                "creation_date": creation_date,
                "status": "Draft",
                "collaborators": collaborators if collaborators else []
            },
            "clauses": []
        }
        
        self.save_contract(contract)
        return contract_id

    def open_contract(self, contract_id):
        contract_path = self._get_contract_path(contract_id)
        if not os.path.exists(contract_path):
            return None
        
        try:
            with self.lock, open(contract_path, "r") as f: # Lock applied to reads
                return json.load(f)
        except (json.JSONDecodeError, FileNotFoundError):
            return None

    def save_contract(self, contract):
        contract_path = self._get_contract_path(contract["metadata"]["contract_id"])
        with self.lock:
            with open(contract_path, "w") as f: # Lock applied to writes
                json.dump(contract, f, indent=4)
                
    def sanitize_filename(self, title):
        '''Remove special characters to make a safe filename'''
        return re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
    
    def convert_to_docx(self, contract_id):
        '''Convert a JSON contract into a DOCX file'''
        contract = self.open_contract(contract_id) 
        if not contract:
            return None, "Contract not found"
        
        doc = Document()
        metadata = contract["metadata"] 
        
        # Use the contract title for the filename
        safe_title = self.sanitize_filename(metadata["title"])
        docx_filename = f"{safe_title}.docx"
        docx_path = os.path.join(self.contract_docx_directory, docx_filename)
        
        # Format creation date
        creation_date = metadata["creation_date"].split("T")[0] # Extract YYYY-MM-DD
        
        # Title
        doc.add_heading(metadata["title"], level=1)
        
        # Contract Info
        doc.add_paragraph(f"Created by: {metadata['creator_name']}")
        doc.add_paragraph(f"Creation Date: {creation_date}")
        doc.add_paragraph(f"Status: {metadata['status']}")
        doc.add_paragraph(f"Description: {metadata['description']}\n")
        
        # Clauses
        # Clause numbering
        clause_number = 1
        for clause in contract["clauses"]:
            doc.add_heading(f"{clause_number}.{clause['short_title']}", level = 3)
            
            # Get latest version of the clause
            latest_version = clause["versions"][0]
            clause_text = latest_version["full_text"]
            
            # Sentence-level numbering
            sentence_number = 1
            for sentence in clause_text.split("\n"):
                sentence = sentence.strip()
                if sentence:
                    doc.add_paragraph(f"{clause_number}.{sentence_number} {sentence}")
                    sentence_number += 1
                
            clause_number += 1 # Increment for the next clause
            
        doc.save(docx_path)
        
        return docx_path, "DOCX file generated successfully"
                
        
        

    def add_clause(self, contract_id, short_title, full_text, publisher):
        contract = self.open_contract(contract_id)
        if not contract:
            return None

        clause_id = self._generate_id()
        new_clause = {
            "clause_id": clause_id,
            "short_title": short_title,
            "versions": [{
                "date": datetime.now().isoformat(),
                "full_text": full_text,
                "publisher": publisher
            }],
            "comments": [] # Initialize empty comments array for this clause
        }
        contract["clauses"].append(new_clause)
        self.save_contract(contract)
        return new_clause

    def update_clause(self, contract_id, clause_id, full_text, publisher_id, publisher_name, short_title=None):
        contract = self.open_contract(contract_id)
        if not contract:
            return False, "Contract not found"

        for clause in contract["clauses"]:
            if clause["clause_id"] == clause_id:
                # Update the clause text
                clause["versions"].insert(0, {
                    "date": datetime.now().isoformat(),
                    "full_text": full_text,
                    "publisher_id": publisher_id,
                    "publisher_name": publisher_name
                })
                
                # Allow renaming if a short_tile is provided
                if short_title:
                    clause["short_title"] = short_title
                    
                self.save_contract(contract)
                return True
        return False
    

    def check_user_permission(self, contract_id, user_id, required_role=None):
        """
        Check if a user has the required role to access a contract.
        Returns True is the user is the creator or has the required role (if specfied).
        """
        contract = self.open_contract(contract_id)
        if not contract:
            return False
        
        if contract["metadata"]["creator_id"] == user_id:
            return True
        
        if required_role:
            for collaborator in contract["metadata"]["collaborators"]:
                if collaborator["user_id"] == user_id and collaborator["role"] == required_role:
                    return True
            return False
        
        # Check if user is a collaborator
        for collaborator in contract["metadata"]["collaborators"]:
            if collaborator["user_id"] == user_id:
                return True
            
        return False
        
    
    def add_collaborator(self, contract_id, collaborator_data, role, added_by):
        """
        Add collaborator with specified role.
        Roles: "Editor", "Viewer:, "Approver"
        Collaborator_data should include user_id, name and email
        """
        contract = self.open_contract(contract_id)
        if not contract:
            return False, "Contract not found"
        
        # Check if user adding collaborator is the creator
        if contract["metadata"]["creator_id"] != added_by:
            return False, "Only the creator can add collaborators"

        # Check if collaborator already exixts
        for collab in contract["metadata"]["collaborators"]:
            if collab["user_id"] == collaborator_data["user_id"]:
                return False, "Collaborator already exists"
        
        # Validate role
        valid_roles = ["Editor", "Viewer", "Approver"]
        if role not in valid_roles:
            return False, "Role must be specified"
        
        # Add collaborator with role
        new_collaborator = {
            "user_id": collaborator_data["user_id"],
            "name": collaborator_data["name"],
            "email": collaborator_data["email"],
            "role": role,
            "added_date": datetime.now().isoformat()
        }
        
        contract["metadata"]["collaborators"].append(new_collaborator)
        self.save_contract(contract)
        return True, "Collaborator added successfully"

    def remove_collaborator(self, contract_id, collaborator_id, removed_by):
        contract = self.open_contract(contract_id)
        if not contract:
            return False
        
        # Check if user removing collaborator is the creator
        if contract["metadata"]["creator_id"] != removed_by:
            return False, "Only the contract creator can remove collaborators"

        # Find and remove collaborator
        for i, collab in enumerate(contract["metadata"]["collaborators"]):
            if collab["user_id"] == collaborator_id:
                contract["metadata"]["collaborators"].pop(i)
                self.save_contract(contract)
                return True, "Collaborator removed successfully"
            
        return False, "Collaborator not found"
    
    def update_role(self, contract_id, collaborator_id, new_role, requester_id):
        contract = self.open_contract(contract_id)
        if not contract:
            return False, "Contract not found"
        
        # Only the contract creator can update roles
        if contract['metadata']['creator_id'] != requester_id:
            return False, "Only the contract creator can update roles"
        
        # Find collaborators and uodate role
        for collab in contract['metadata']['collaborators']:
            if collab['user_id'] == collaborator_id:
                collab['role'] = new_role
                self.save_contract(contract)
                return True, "Role updated successfully"
        return False, "Collaborator not found"

    # Get all clauses
    def get_clauses(self, contract_id):
        contract = self.open_contract(contract_id)
        if not contract:
            return False

        if "clauses" in contract:
            return contract["clauses"]
        
        return None
        
         
        
    def delete_clause(self, contract_id, clause_id):
        contract = self.open_contract(contract_id)
        if not contract:
            return False
        
        contract["clauses"] = [clause for clause in contract["clauses"] if clause["clause_id"] != clause_id]
        self.save_contract(contract)
        return True

    def delete_contract(self, contract_id):
        contract_path = self._get_contract_path(contract_id)
        if os.path.exists(contract_path):
            with self.lock:
                os.remove(contract_path)
            return True
        return False

    def list_contracts(self, creator_id=None, collaborator_id=None):
        contracts = []
        for filename in os.listdir(self.contract_directory):
            if filename.endswith(".json"):
                contract_id = filename[:-5]
                contract = self.open_contract(contract_id)
                if contract:
                    if creator_id and contract["metadata"]["creator_id"] != creator_id:
                        continue
                    if collaborator_id and collaborator_id not in contract["metadata"]["collaborators"]:
                        continue
                    contracts.append(contract["metadata"])
        return contracts
    
    def add_comment(self, contract_id, clause_id, user_id, email, name, comment_text):
        """
        Add a comment to a specific clause in a contract.
        Any user with access to the contract can comment.
        """
            
        contract = self.open_contract(contract_id)
        if not contract:
            return False, "Contract not found"
            
            # Check if user has access
        has_access = False
        if contract["metadata"]["creator_id"] == user_id:
            has_access = True
        else:
            for collab in contract["metadata"]["collaborators"]:
                if collab["user_id"] == user_id:
                    has_access = True
                    break
                    
        if not has_access:
            return False, "User does not have access to this contract"
            
        # Find the clause
        for clause in contract["clauses"]:
            if clause["clause_id"] == clause_id:
                # Initialize comments array if it doesn't exist
                if "comments" not in clause:
                    clause["comments"] = []
                        
                # Add the comment
                comment_id = self._generate_id()
                new_comment = {
                        "comment_id": comment_id,
                        "user_id": user_id,
                        "email": email,
                        "name": name,
                        "comment": comment_text,
                        "date": datetime.now().isoformat()
                    }
                    
                clause["comments"].append(new_comment)
                self.save_contract(contract)
                return True, comment_id
        return False, "Clause not found"
    
    def get_comments(self, contract_id, clause_id):
        """
        Get all comments for a specific clause
        """
        contract = self.open_contract(contract_id)
        if not contract:
            return None
        
        for clause in contract["clauses"]:
            if clause["clause_id"] == clause_id:
                if "comments" in clause:
                    return clause["comments"]
                return []
        
        return None
    
    def delete_comment(self, contract_id, clause_id, comment_id, user_id):
        """
        Delete a comment. Only the comment creator or contract creator can delete.
        """
        
        contract = self.open_contract(contract_id)
        if not contract:
            return False, "Contract not found"
        
        for clause in contract["clauses"]:
            if clause["clause_id"] == clause_id and "comments" in clause:
                for i, comment in enumerate(clause["comments"]):
                    if comment["comment_id"] == comment_id:
                        # Check if user is authorized to delete
                        if comment['user_id'] == user_id or contract["metadata"]["creator_id"] == user_id:
                            clause["comments"].pop(i)
                            self.save_contract(contract)
                            return True, "Comment deleted successfully"
                        else: 
                            return False, "Not authorized to delete this comment"
                return False, "Comment not found"
        return False, "Clause not found"
    
    def move_clause(self, contract_id, clause_id, new_index):
        contract = self.open_contract(contract_id)
        if not contract:
            return False, "Contract not found"
        
        clauses = contract["clauses"]
        
        # Find the clause to move
        clause_to_move = next((clause for clause in clauses if clause["clause_id"] == clause_id), None)
        if not clause_to_move:
            return False, "Clause not found"
        
        # Remove it from its current position
        clauses.remove(clause_to_move)
        
        # Insert it into the new position
        if new_index < 0 or new_index >= len(clauses):
            clauses.append(clause_to_move) # Move to end if index is out of bounds
        else: 
            clauses.insert(new_index, clause_to_move)
        
        # Save updated contract
        self.save_contract(contract)
        return True, "Clause moved successfully"
    
    def approve_contract(self, contract_id, user_id):
        contract = self.open_contract(contract_id)
        if not contract:
            return False, "Contract not found"

        # Only Approvers can approve the contract
        has_permission = any(
            collab["user_id"] == user_id and collab['role'] == 'Approver'
            for collab in contract['metadata']['collaborators']
        )      
        if not has_permission:
            return False, "Only Approvers can approve the contract" 
        
        # Change status to Approved
        contract['metadata']['status'] = 'Approved'
        self.save_contract(contract)
        
        return True, "Contract approved successfully"
    
    def explain_clause(self, clause_text):
        """Use OpenAI API to explain a contract clause in simple terms."""
        prompt = f"{clause_text}"
        
        response = client.chat.completions.create(
            model ="gpt-4",
            messages=[{"role": "system", "content": "You are a legal AI assistant that explains contract clauses. Explain the provided clause from a contract in a simple and clear way, in not more than 200 words:"},
                      {"role": "user", "content": prompt}]
        )
        
        return response.choices[0].message.content.strip()
                     